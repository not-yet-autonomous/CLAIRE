# Copyright (c) 2026 James Cole. Licensed under the MIT License.
"""
CLAIRE — Build 8: Output Layer
Input:   data/candidates_track_a/b/c.json
Output:  output/CLAIRE_Weekly_Digest_YYYY-MM-DD.docx  (default)
         output/claire_digest_YYYY-MM-DD.pdf           (--format pdf)
         skill_drafts/*.md (Track B skill skeletons)

Run:     python claire_output.py
         python claire_output.py --date 2026-04-25  (override date)
         python claire_output.py --format pdf
"""

import json
import sys
import argparse
import logging
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

# reportlab — PDF generation (Build 8)
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak, KeepTogether,
)
from reportlab.platypus.flowables import HRFlowable

load_dotenv()

# ─────────────────────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────────────────────

BASE_DIR        = Path(__file__).parent
DATA_DIR        = BASE_DIR / "data"
OUTPUT_DIR      = BASE_DIR / "output"
SKILL_DRAFT_DIR      = BASE_DIR / "skill_drafts"

# CLAIRE-A shadow output paths
CLAIRE_A_DATA_DIR    = BASE_DIR / "data"
LOGS_DIR        = BASE_DIR / "logs"

OUTPUT_DIR.mkdir(exist_ok=True)
SKILL_DRAFT_DIR.mkdir(exist_ok=True)
LOGS_DIR.mkdir(exist_ok=True)

OUTPUT_LOG_PATH = LOGS_DIR / "output.log"

CANDIDATE_PATHS = {
    "a": DATA_DIR / "candidates_track_a.json",
    "b": DATA_DIR / "candidates_track_b.json",
    "c": DATA_DIR / "candidates_track_c.json",
}

TRIAGE_TAGGED_PATH = DATA_DIR / "tagged_posts.json"

# Waterton brand colors
NAVY        = RGBColor(0x1E, 0x3A, 0x5F)
GOLD        = RGBColor(0xC5, 0xA0, 0x30)
CHARCOAL    = RGBColor(0x3C, 0x3C, 0x3C)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY  = "F7F8FA"
BORDER_GRAY = "E2E8F0"
NAVY_HEX    = "1E3A5F"
GOLD_HEX    = "C5A030"
GREEN_HEX   = "2D6A4F"
AMBER_HEX   = "8B6914"

# ─────────────────────────────────────────────────────────────────────────────
# LOGGING
# ─────────────────────────────────────────────────────────────────────────────

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%H:%M:%S",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler(OUTPUT_LOG_PATH, encoding="utf-8"),
    ]
)
log = logging.getLogger("claire.output")

# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT HELPERS — Waterton theme
# ─────────────────────────────────────────────────────────────────────────────

def init_document() -> Document:
    """Create and configure a Waterton-themed document."""
    doc = Document()

    # US Letter, 1" margins
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name      = "Century Gothic"
    normal.font.size      = Pt(11)
    normal.font.color.rgb = CHARCOAL
    normal.paragraph_format.space_after  = Pt(6)

    # Heading styles
    def set_heading(style_name, size, color, before, after):
        s = doc.styles[style_name]
        s.font.name      = "Georgia"
        s.font.size      = Pt(size)
        s.font.bold      = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after  = Pt(after)

    set_heading("Heading 1", 22, NAVY,  18, 6)
    set_heading("Heading 2", 16, GOLD,  12, 4)
    set_heading("Heading 3", 13, GOLD,  10, 3)

    # Strip bottom borders from Title and Heading styles in default template.
    # Word's default.dotx embeds a blue bottom border on the Title style
    # which renders as a line through large paragraph text.
    def clear_style_borders(style):
        pPr = style.element.find(qn("w:pPr"))
        if pPr is not None:
            for pBdr in pPr.findall(qn("w:pBdr")):
                pPr.remove(pBdr)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            clear_style_borders(doc.styles[style_name])
        except KeyError:
            pass

    return doc


def clear_paragraph_borders(para):
    """Explicitly remove all paragraph borders inherited from template."""
    pPr = para._p.get_or_add_pPr()
    for pBdr in pPr.findall(qn("w:pBdr")):
        pPr.remove(pBdr)
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right", "between"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        pBdr.append(el)
    pPr.append(pBdr)


def add_title_block(doc: Document, date_str: str):
    """Add CLAIRE document title block."""
    # Eyebrow
    eyebrow = doc.add_paragraph()
    run = eyebrow.add_run("Weekly Intelligence Digest")
    run.font.name      = "Georgia"
    run.font.size      = Pt(11)
    run.font.italic    = True
    run.font.color.rgb = GOLD
    eyebrow.paragraph_format.space_before = Pt(0)
    eyebrow.paragraph_format.space_after  = Pt(4)
    eyebrow.paragraph_format.line_spacing = None
    clear_paragraph_borders(eyebrow)

    # Title
    title_para = doc.add_paragraph()
    run = title_para.add_run("CLAIRE")
    run.font.name      = "Georgia"
    run.font.size      = Pt(28)
    run.font.bold      = True
    run.font.color.rgb = NAVY
    title_para.paragraph_format.space_before = Pt(6)
    title_para.paragraph_format.space_after  = Pt(6)
    title_para.paragraph_format.line_spacing = None
    clear_paragraph_borders(title_para)

    # Subtitle
    sub = doc.add_paragraph()
    run = sub.add_run(f"Claude Configuration Signal Report — {date_str}")
    run.font.name      = "Georgia"
    run.font.size      = Pt(14)
    run.font.bold      = True
    run.font.color.rgb = GOLD
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after  = Pt(28)
    sub.paragraph_format.line_spacing = None
    clear_paragraph_borders(sub)


def add_rule(doc: Document, color_hex: str = GOLD_HEX):
    """Add a horizontal rule."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "8")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def shade_cell(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def set_cell_borders(cell, color: str = BORDER_GRAY):
    tcPr  = cell._tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:color"), color)
        tcBdr.append(el)
    tcPr.append(tcBdr)


def header_cell(cell, text: str):
    shade_cell(cell, NAVY_HEX)
    set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = cell.paragraphs[0].add_run(text)
    run.font.name      = "Century Gothic"
    run.font.size      = Pt(10)
    run.font.bold      = True
    run.font.color.rgb = WHITE


def body_cell(cell, text: str, even: bool = True):
    shade_cell(cell, "FFFFFF" if even else LIGHT_GRAY)
    set_cell_borders(cell)
    run = cell.paragraphs[0].add_run(str(text))
    run.font.name      = "Century Gothic"
    run.font.size      = Pt(10)
    run.font.color.rgb = CHARCOAL


def add_confidence_badge(para, confidence: str):
    """Inline confidence indicator."""
    run = para.add_run(f" [{confidence}]")
    run.font.name = "Century Gothic"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = NAVY if confidence == "HIGH" else GOLD


def add_callout(doc: Document, text: str):
    """Light gray callout box with gold left border — for hypothesis prompt."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:color"), GOLD_HEX)
    left.set(qn("w:space"), "20")
    pBdr.append(left)
    pPr.append(pBdr)
    # Background
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  LIGHT_GRAY)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name      = "Century Gothic"
    run.font.size      = Pt(10)
    run.font.italic    = True
    run.font.color.rgb = CHARCOAL


# ─────────────────────────────────────────────────────────────────────────────
# SECTION BUILDERS
# ─────────────────────────────────────────────────────────────────────────────

def build_signal_summary(doc: Document, triage_data: dict):
    """Section 1 — Signal Summary."""
    doc.add_heading("Signal Summary", level=1)

    meta  = triage_data.get("meta", {})
    stats = meta.get("stats", {})

    # Run metadata
    run_at = meta.get("run_at", "unknown")
    total  = meta.get("total", 0)

    para = doc.add_paragraph()
    para.add_run(f"Run date: ").bold = True
    para.add_run(run_at[:10])
    para2 = doc.add_paragraph()
    para2.add_run("Posts scanned: ").bold = True
    para2.add_run(str(total))

    doc.add_heading("Signal Breakdown", level=2)

    # Signal type table
    by_signal  = stats.get("by_signal", {})
    by_persona = stats.get("by_persona", {})
    by_source  = stats.get("by_source", {})

    headers = ["Category", "Count"]
    rows = [
        ["behavior_complaint", str(by_signal.get("behavior_complaint", 0))],
        ["workflow_gap",       str(by_signal.get("workflow_gap", 0))],
        ["feature_praise",     str(by_signal.get("feature_praise", 0))],
        ["competitor_gap",     str(by_signal.get("competitor_gap", 0))],
        ["cross_platform",     str(by_signal.get("cross_platform_workflow", 0))],
        ["noise (dropped)",    str(by_signal.get("noise", 0))],
    ]

    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_widths = [6240, 3120]
    for i, col in enumerate(table.columns):
        col.width = col_widths[i]

    for i, h in enumerate(headers):
        header_cell(table.rows[0].cells[i], h)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            body_cell(table.rows[r_idx + 1].cells[c_idx], val, even=(r_idx % 2 == 0))

    doc.add_paragraph()
    doc.add_heading("Developer Signal Disposition", level=2)
    dev_count = by_persona.get("developer", 0)
    p = doc.add_paragraph()
    p.add_run(f"{dev_count} developer-persona posts filtered at cross-reference gate "
              f"per locked pipeline configuration (filter_out).")

    add_rule(doc)


def build_track_a_section(doc: Document, candidates: dict):
    """Section 2 — Memory Edit Candidates and Profile Diffs."""
    doc.add_heading("Track A — Claude Native Signal", level=1)

    memory_edits  = candidates.get("memory_edit_candidates", [])
    profile_diffs = candidates.get("profile_diff_candidates", [])
    watch_list    = candidates.get("behavior_watch", [])

    # Memory edit candidates
    doc.add_heading("Memory Edit Candidates", level=2)
    if not memory_edits:
        doc.add_paragraph("No memory edit candidates this cycle. "
                          "Signal below evidence threshold or no patterns identified.")
    else:
        for i, candidate in enumerate(memory_edits, 1):
            para = doc.add_paragraph()
            run = para.add_run(f"{i}. {candidate.get('control', '')}")
            run.font.bold      = True
            run.font.name      = "Century Gothic"
            run.font.color.rgb = NAVY
            add_confidence_badge(para, candidate.get("confidence", "MEDIUM"))

            doc.add_paragraph(candidate.get("rationale", "")).style = doc.styles["Normal"]

            # Ready-to-apply command
            add_callout(doc,
                f"Apply command: memory_user_edits → add → \"{candidate.get('control', '')}\"\n"
                f"Write your hypothesis before applying.")

            # Source links
            sources = candidate.get("source_posts", [])
            if sources:
                src_para = doc.add_paragraph()
                src_para.add_run("Sources: ").bold = True
                src_para.add_run(" · ".join(
                    f"reddit.com{p}" if p.startswith("/r/") else p
                    for p in sources[:5]
                ))
                src_para.runs[-1].font.size = Pt(9)

            if i < len(memory_edits):
                doc.add_paragraph()

    add_rule(doc)

    # Profile diff candidates
    doc.add_heading("Profile Diff Candidates", level=2)
    if not profile_diffs:
        doc.add_paragraph("No profile diff candidates this cycle.")
    else:
        for i, candidate in enumerate(profile_diffs, 1):
            para = doc.add_paragraph()
            run = para.add_run(
                f"{i}. Target section: {candidate.get('target_section', '')}")
            run.font.bold      = True
            run.font.name      = "Century Gothic"
            run.font.color.rgb = NAVY
            add_confidence_badge(para, candidate.get("confidence", "MEDIUM"))

            doc.add_paragraph(
                f"Proposed change: {candidate.get('proposed_change', '')}"
            ).style = doc.styles["Normal"]
            doc.add_paragraph(candidate.get("rationale", "")).style = doc.styles["Normal"]

            add_callout(doc, "Review proposed change against current profile section "
                            "before applying. Write hypothesis first.")

            sources = candidate.get("source_posts", [])
            if sources:
                src_para = doc.add_paragraph()
                src_para.add_run("Sources: ").bold = True
                src_para.add_run(" · ".join(sources[:5]))
                src_para.runs[-1].font.size = Pt(9)

    add_rule(doc)

    # Behavior watch list
    doc.add_heading("Behavior Watch List", level=2)
    doc.add_paragraph(
        "Patterns flagged for monitoring — not yet actionable. "
        "Check against personal session notes before next cycle."
    )
    if not watch_list:
        doc.add_paragraph("Nothing on watch list this cycle.")
    else:
        for item in watch_list:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(item.get("pattern", "")).bold = True
            doc.add_paragraph(
                f"Why not actioned: {item.get('why_not_actioned', '')}"
            ).style = doc.styles["Normal"]

    add_rule(doc)


def build_track_b_section(doc: Document, candidates: dict):
    """Section 3 — Competitor Gap Signal."""
    doc.add_heading("Track B — Competitor Gap Signal", level=1)

    skill_drafts     = candidates.get("skill_draft_candidates", [])
    profile_additions = candidates.get("profile_addition_candidates", [])

    doc.add_heading("Skill Draft Candidates", level=2)
    if not skill_drafts:
        doc.add_paragraph(
            "No skill draft candidates this cycle. "
            "Competitor gap signal below evidence threshold of 3 corroborating posts. "
            "This is expected in early cycles — pattern builds over weekly runs."
        )
    else:
        for i, candidate in enumerate(skill_drafts, 1):
            para = doc.add_paragraph()
            run = para.add_run(f"{i}. {candidate.get('skill_name', '')}")
            run.font.bold      = True
            run.font.color.rgb = NAVY
            run.font.name      = "Century Gothic"
            add_confidence_badge(para, candidate.get("confidence", "MEDIUM"))

            doc.add_paragraph(
                f"Gap: {candidate.get('gap_description', '')}"
            ).style = doc.styles["Normal"]
            doc.add_paragraph(
                f"Trigger: {candidate.get('trigger_description', '')}"
            ).style = doc.styles["Normal"]
            doc.add_paragraph(
                f"Estimated effort: {candidate.get('estimated_build_effort', '')}"
            ).style = doc.styles["Normal"]

            add_callout(doc,
                f"SKILL.md skeleton saved to skill_drafts/{candidate.get('skill_name', 'skill').replace(' ', '_')}.md")

    add_rule(doc)

    doc.add_heading("Profile Addition Candidates", level=2)
    if not profile_additions:
        doc.add_paragraph("No profile addition candidates this cycle.")
    else:
        for i, candidate in enumerate(profile_additions, 1):
            para = doc.add_paragraph()
            run = para.add_run(f"{i}. Section: {candidate.get('target_section', '')}")
            run.font.bold      = True
            run.font.color.rgb = NAVY
            run.font.name      = "Century Gothic"
            add_confidence_badge(para, candidate.get("confidence", "MEDIUM"))

            doc.add_paragraph(candidate.get("proposed_text", "")).style = doc.styles["Normal"]
            doc.add_paragraph(candidate.get("rationale", "")).style = doc.styles["Normal"]

    add_rule(doc)


def build_track_c_section(doc: Document, candidates: dict):
    """Section 4 — Cross-Platform Workflow Techniques."""
    doc.add_heading("Track C — Cross-Platform Workflow Techniques", level=1)
    doc.add_paragraph(
        "Portable techniques to test manually before any configuration change is considered."
    )

    techniques = candidates.get("technique_candidates", [])

    if not techniques:
        doc.add_paragraph("No technique candidates this cycle.")
    else:
        for i, technique in enumerate(techniques, 1):
            para = doc.add_paragraph()
            run = para.add_run(f"{i}. {technique.get('technique_name', '')}")
            run.font.bold      = True
            run.font.color.rgb = NAVY
            run.font.name      = "Century Gothic"
            add_confidence_badge(para, technique.get("confidence", "MEDIUM"))

            doc.add_paragraph(technique.get("description", "")).style = doc.styles["Normal"]
            add_callout(doc, f"Test: {technique.get('test_suggestion', '')}")

    add_rule(doc)


def build_eval_section(doc: Document, date_str: str):
    """Section 5 — Eval Status."""
    doc.add_heading("Eval Loop Status", level=1)

    doc.add_heading("Applied Changes Pending Eval", level=2)
    doc.add_paragraph(
        "Review change_log.json for changes applied since last digest. "
        "For each change where a relevant session has occurred, add an eval note:"
    )
    add_callout(doc,
        f"{date_str} | [session context] | held/partial/no | [one sentence notes]")

    doc.add_heading("Friction Log Reminder", level=2)
    doc.add_paragraph(
        "Update friction_log.txt (project root — not data/) this week. "
        "2-4 entries. Specific behaviors in specific contexts. "
        "Blank weeks score everything MEDIUM next cycle — the cross-reference gate loses precision."
    )

    doc.add_heading("Quarterly Eval", level=2)
    doc.add_paragraph(
        "If quarterly eval is due: feed change_log.json and friction_log.txt to Claude. "
        "Request eval report. Revert changes that did not hold. "
        "Flag source signal from reverted changes as low-quality in archive."
    )


# ─────────────────────────────────────────────────────────────────────────────
# SKILL DRAFT FILE WRITER
# ─────────────────────────────────────────────────────────────────────────────

def write_skill_drafts(candidates_b: dict) -> int:
    """Write SKILL.md skeleton files for Track B skill candidates."""
    skill_drafts = candidates_b.get("skill_draft_candidates", [])
    written = 0

    for candidate in skill_drafts:
        name     = candidate.get("skill_name", "unnamed_skill")
        skeleton = candidate.get("skill_md_skeleton", "")
        filename = name.lower().replace(" ", "_").replace("/", "_") + ".md"
        path     = SKILL_DRAFT_DIR / filename

        content = f"""---
name: {name}
description: {candidate.get('trigger_description', '')}
status: DRAFT — generated by CLAIRE, requires human review and implementation
generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}
confidence: {candidate.get('confidence', 'MEDIUM')}
estimated_effort: {candidate.get('estimated_build_effort', 'unknown')}
source_signal: {', '.join(candidate.get('source_posts', [])[:3])}
---

# {name}

## Gap Description
{candidate.get('gap_description', '')}

## Trigger
{candidate.get('trigger_description', '')}

## Skeleton
{skeleton}

## Implementation Notes
(Complete this section before installing the skill)

## Review Status
- [ ] Gap confirmed from personal session experience
- [ ] Hypothesis written
- [ ] Implementation complete
- [ ] Tested in session
- [ ] Installed
"""
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        log.info(f"Skill draft written → {filename}")
        written += 1

    return written


# ─────────────────────────────────────────────────────────────────────────────
# PDF GENERATION — reportlab (Build 8)
# Font note: reportlab's standard 14 fonts do not include Century Gothic or
# Georgia. Helvetica is used in place of Century Gothic; Times-Roman in place
# of Georgia. This is a deliberate fallback — registering TTF fonts would
# require the font files to be present at runtime and is not worth the
# complexity for an internal digest.
#
# Callout boxes: replicated as a 1-column Table with a thick left border line
# (gold) drawn via TableStyle LINEABOVE/LINEBELOW + a Paragraph. ReportLab has
# no native sidebar-border primitive for Platypus flowables; the Table approach
# is the standard pattern.
# ─────────────────────────────────────────────────────────────────────────────

# PDF brand colors
PDF_NAVY    = HexColor("#1E3A5F")
PDF_GOLD    = HexColor("#C5A030")
PDF_CHARCOAL = HexColor("#3C3C3C")
PDF_LGRAY   = HexColor("#F7F8FA")
PDF_BGRAY   = HexColor("#E2E8F0")
PDF_WHITE   = colors.white
PDF_GREEN   = HexColor("#2D6A4F")
PDF_AMBER   = HexColor("#8B6914")


def build_pdf_styles() -> dict:
    """Return a dict of named ParagraphStyle objects for the digest."""
    base = getSampleStyleSheet()

    styles = {}

    styles["normal"] = ParagraphStyle(
        "pdf_normal",
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        textColor=PDF_CHARCOAL,
        spaceAfter=6,
    )
    styles["eyebrow"] = ParagraphStyle(
        "pdf_eyebrow",
        fontName="Times-Italic",
        fontSize=11,
        leading=14,
        textColor=PDF_GOLD,
        spaceAfter=2,
    )
    styles["title"] = ParagraphStyle(
        "pdf_title",
        fontName="Times-Bold",
        fontSize=28,
        leading=32,
        textColor=PDF_NAVY,
        spaceBefore=4,
        spaceAfter=4,
    )
    styles["subtitle"] = ParagraphStyle(
        "pdf_subtitle",
        fontName="Times-Bold",
        fontSize=14,
        leading=18,
        textColor=PDF_GOLD,
        spaceAfter=18,
    )
    styles["h1"] = ParagraphStyle(
        "pdf_h1",
        fontName="Times-Bold",
        fontSize=16,
        leading=20,
        textColor=PDF_NAVY,
        spaceBefore=14,
        spaceAfter=6,
    )
    styles["h2"] = ParagraphStyle(
        "pdf_h2",
        fontName="Times-Bold",
        fontSize=13,
        leading=17,
        textColor=PDF_GOLD,
        spaceBefore=10,
        spaceAfter=4,
    )
    styles["h3"] = ParagraphStyle(
        "pdf_h3",
        fontName="Times-Bold",
        fontSize=11,
        leading=14,
        textColor=PDF_GOLD,
        spaceBefore=8,
        spaceAfter=3,
    )
    styles["body_bold"] = ParagraphStyle(
        "pdf_body_bold",
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=14,
        textColor=PDF_NAVY,
        spaceAfter=2,
    )
    styles["small"] = ParagraphStyle(
        "pdf_small",
        fontName="Helvetica",
        fontSize=8,
        leading=11,
        textColor=PDF_CHARCOAL,
        spaceAfter=4,
    )
    styles["callout"] = ParagraphStyle(
        "pdf_callout",
        fontName="Helvetica-Oblique",
        fontSize=9,
        leading=13,
        textColor=PDF_CHARCOAL,
        leftIndent=8,
        rightIndent=8,
        spaceBefore=2,
        spaceAfter=2,
        backColor=PDF_LGRAY,
    )
    styles["bullet"] = ParagraphStyle(
        "pdf_bullet",
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        textColor=PDF_CHARCOAL,
        leftIndent=12,
        bulletIndent=0,
        spaceAfter=4,
    )
    styles["table_header"] = ParagraphStyle(
        "pdf_table_header",
        fontName="Helvetica-Bold",
        fontSize=9,
        leading=12,
        textColor=PDF_WHITE,
    )
    styles["table_body"] = ParagraphStyle(
        "pdf_table_body",
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=PDF_CHARCOAL,
    )
    return styles


def pdf_rule(styles: dict) -> HRFlowable:
    """Gold horizontal rule."""
    return HRFlowable(
        width="100%",
        thickness=2,
        color=PDF_GOLD,
        spaceAfter=8,
        spaceBefore=4,
    )


def pdf_callout(text: str, styles: dict) -> Table:
    """Gold left-bordered callout box.

    ReportLab Platypus has no sidebar-border primitive for arbitrary
    flowables. The standard pattern is a 1-cell Table with a thick left
    LINEBEFORE rule in TableStyle. The background is set via BACKGROUND.
    """
    cell_para = Paragraph(text.replace("\n", "<br/>"), styles["callout"])
    t = Table([[cell_para]], colWidths=["100%"])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), PDF_LGRAY),
        ("LINEBEFORE",  (0, 0), (0, -1), 4, PDF_GOLD),
        ("TOPPADDING",  (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING",   (0, 0), (-1, -1), 10),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 8),
    ]))
    return t


def pdf_confidence_label(confidence: str) -> str:
    """Return inline XML markup for confidence badge."""
    color = "#1E3A5F" if confidence == "HIGH" else "#C5A030"
    return f' <font color="{color}" size="8"><b>[{confidence}]</b></font>'


def pdf_build_signal_summary(story: list, triage_data: dict, styles: dict):
    """PDF Section 1 — Signal Summary."""
    story.append(Paragraph("Signal Summary", styles["h1"]))

    meta  = triage_data.get("meta", {})
    stats = meta.get("stats", {})
    run_at = meta.get("run_at", "unknown")
    total  = meta.get("total", 0)

    story.append(Paragraph(
        f'<b>Run date:</b> {run_at[:10]}', styles["normal"]))
    story.append(Paragraph(
        f'<b>Posts scanned:</b> {total}', styles["normal"]))

    story.append(Paragraph("Signal Breakdown", styles["h2"]))

    by_signal  = stats.get("by_signal", {})
    by_persona = stats.get("by_persona", {})

    headers = [
        Paragraph("Category", styles["table_header"]),
        Paragraph("Count",    styles["table_header"]),
    ]
    rows = [
        ["behavior_complaint", str(by_signal.get("behavior_complaint", 0))],
        ["workflow_gap",       str(by_signal.get("workflow_gap", 0))],
        ["feature_praise",     str(by_signal.get("feature_praise", 0))],
        ["competitor_gap",     str(by_signal.get("competitor_gap", 0))],
        ["cross_platform",     str(by_signal.get("cross_platform_workflow", 0))],
        ["noise (dropped)",    str(by_signal.get("noise", 0))],
    ]

    table_data = [headers]
    for i, (cat, cnt) in enumerate(rows):
        bg = PDF_WHITE if i % 2 == 0 else PDF_LGRAY
        table_data.append([
            Paragraph(cat, styles["table_body"]),
            Paragraph(cnt, styles["table_body"]),
        ])

    t = Table(table_data, colWidths=[4.5 * inch, 2.0 * inch])
    ts = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), PDF_NAVY),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [PDF_WHITE, PDF_LGRAY]),
        ("GRID",       (0, 0), (-1, -1), 0.5, PDF_BGRAY),
        ("TOPPADDING",    (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING",   (0, 0), (-1, -1), 6),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
    ])
    t.setStyle(ts)
    story.append(t)
    story.append(Spacer(1, 8))

    story.append(Paragraph("Developer Signal Disposition", styles["h2"]))
    dev_count = by_persona.get("developer", 0)
    story.append(Paragraph(
        f"{dev_count} developer-persona posts filtered at cross-reference gate "
        f"per locked pipeline configuration (filter_out).", styles["normal"]))

    story.append(pdf_rule(styles))


def pdf_build_track_a_section(story: list, candidates: dict, styles: dict):
    """PDF Section 2 — Track A."""
    story.append(Paragraph("Track A — Claude Native Signal", styles["h1"]))

    memory_edits  = candidates.get("memory_edit_candidates", [])
    profile_diffs = candidates.get("profile_diff_candidates", [])
    watch_list    = candidates.get("behavior_watch", [])

    story.append(Paragraph("Memory Edit Candidates", styles["h2"]))
    if not memory_edits:
        story.append(Paragraph(
            "No memory edit candidates this cycle. "
            "Signal below evidence threshold or no patterns identified.",
            styles["normal"]))
    else:
        for i, candidate in enumerate(memory_edits, 1):
            conf = candidate.get("confidence", "MEDIUM")
            label = pdf_confidence_label(conf)
            story.append(Paragraph(
                f'<b>{i}. {candidate.get("control", "")}</b>{label}',
                styles["body_bold"]))
            story.append(Paragraph(candidate.get("rationale", ""), styles["normal"]))
            story.append(pdf_callout(
                f'Apply command: memory_user_edits → add → '
                f'"{candidate.get("control", "")}"\n'
                f'Write your hypothesis before applying.',
                styles))
            sources = candidate.get("source_posts", [])
            if sources:
                src_text = " · ".join(
                    f"reddit.com{p}" if p.startswith("/r/") else p
                    for p in sources[:5]
                )
                story.append(Paragraph(
                    f'<b>Sources:</b> {src_text}', styles["small"]))
            if i < len(memory_edits):
                story.append(Spacer(1, 6))

    story.append(pdf_rule(styles))

    story.append(Paragraph("Profile Diff Candidates", styles["h2"]))
    if not profile_diffs:
        story.append(Paragraph(
            "No profile diff candidates this cycle.", styles["normal"]))
    else:
        for i, candidate in enumerate(profile_diffs, 1):
            conf = candidate.get("confidence", "MEDIUM")
            label = pdf_confidence_label(conf)
            story.append(Paragraph(
                f'<b>{i}. Target section: {candidate.get("target_section", "")}</b>{label}',
                styles["body_bold"]))
            story.append(Paragraph(
                f'Proposed change: {candidate.get("proposed_change", "")}',
                styles["normal"]))
            story.append(Paragraph(candidate.get("rationale", ""), styles["normal"]))
            story.append(pdf_callout(
                "Review proposed change against current profile section "
                "before applying. Write hypothesis first.",
                styles))
            sources = candidate.get("source_posts", [])
            if sources:
                story.append(Paragraph(
                    f'<b>Sources:</b> {" · ".join(sources[:5])}',
                    styles["small"]))

    story.append(pdf_rule(styles))

    story.append(Paragraph("Behavior Watch List", styles["h2"]))
    story.append(Paragraph(
        "Patterns flagged for monitoring — not yet actionable. "
        "Check against personal session notes before next cycle.",
        styles["normal"]))
    if not watch_list:
        story.append(Paragraph("Nothing on watch list this cycle.", styles["normal"]))
    else:
        for item in watch_list:
            story.append(Paragraph(
                f'• <b>{item.get("pattern", "")}</b>',
                styles["bullet"]))
            story.append(Paragraph(
                f'Why not actioned: {item.get("why_not_actioned", "")}',
                styles["normal"]))

    story.append(pdf_rule(styles))


def pdf_build_track_b_section(story: list, candidates: dict, styles: dict):
    """PDF Section 3 — Track B."""
    story.append(Paragraph("Track B — Competitor Gap Signal", styles["h1"]))

    skill_drafts      = candidates.get("skill_draft_candidates", [])
    profile_additions = candidates.get("profile_addition_candidates", [])

    story.append(Paragraph("Skill Draft Candidates", styles["h2"]))
    if not skill_drafts:
        story.append(Paragraph(
            "No skill draft candidates this cycle. "
            "Competitor gap signal below evidence threshold of 3 corroborating posts. "
            "This is expected in early cycles — pattern builds over weekly runs.",
            styles["normal"]))
    else:
        for i, candidate in enumerate(skill_drafts, 1):
            conf  = candidate.get("confidence", "MEDIUM")
            label = pdf_confidence_label(conf)
            story.append(Paragraph(
                f'<b>{i}. {candidate.get("skill_name", "")}</b>{label}',
                styles["body_bold"]))
            story.append(Paragraph(
                f'Gap: {candidate.get("gap_description", "")}',
                styles["normal"]))
            story.append(Paragraph(
                f'Trigger: {candidate.get("trigger_description", "")}',
                styles["normal"]))
            story.append(Paragraph(
                f'Estimated effort: {candidate.get("estimated_build_effort", "")}',
                styles["normal"]))
            sname = candidate.get("skill_name", "skill").replace(" ", "_")
            story.append(pdf_callout(
                f"SKILL.md skeleton saved to skill_drafts/{sname}.md",
                styles))

    story.append(pdf_rule(styles))

    story.append(Paragraph("Profile Addition Candidates", styles["h2"]))
    if not profile_additions:
        story.append(Paragraph(
            "No profile addition candidates this cycle.", styles["normal"]))
    else:
        for i, candidate in enumerate(profile_additions, 1):
            conf  = candidate.get("confidence", "MEDIUM")
            label = pdf_confidence_label(conf)
            story.append(Paragraph(
                f'<b>{i}. Section: {candidate.get("target_section", "")}</b>{label}',
                styles["body_bold"]))
            story.append(Paragraph(candidate.get("proposed_text", ""), styles["normal"]))
            story.append(Paragraph(candidate.get("rationale", ""), styles["normal"]))

    story.append(pdf_rule(styles))


def pdf_build_track_c_section(story: list, candidates: dict, styles: dict):
    """PDF Section 4 — Track C."""
    story.append(Paragraph(
        "Track C — Cross-Platform Workflow Techniques", styles["h1"]))
    story.append(Paragraph(
        "Portable techniques to test manually before any configuration change is considered.",
        styles["normal"]))

    techniques = candidates.get("technique_candidates", [])

    if not techniques:
        story.append(Paragraph("No technique candidates this cycle.", styles["normal"]))
    else:
        for i, technique in enumerate(techniques, 1):
            conf  = technique.get("confidence", "MEDIUM")
            label = pdf_confidence_label(conf)
            story.append(Paragraph(
                f'<b>{i}. {technique.get("technique_name", "")}</b>{label}',
                styles["body_bold"]))
            story.append(Paragraph(technique.get("description", ""), styles["normal"]))
            story.append(pdf_callout(
                f'Test: {technique.get("test_suggestion", "")}',
                styles))

    story.append(pdf_rule(styles))


def pdf_build_eval_section(story: list, date_str: str, styles: dict):
    """PDF Section 5 — Eval Status."""
    story.append(Paragraph("Eval Loop Status", styles["h1"]))

    story.append(Paragraph("Applied Changes Pending Eval", styles["h2"]))
    story.append(Paragraph(
        "Review change_log.json for changes applied since last digest. "
        "For each change where a relevant session has occurred, add an eval note:",
        styles["normal"]))
    story.append(pdf_callout(
        f"{date_str} | [session context] | held/partial/no | [one sentence notes]",
        styles))

    story.append(Paragraph("Friction Log Reminder", styles["h2"]))
    story.append(Paragraph(
        "Update friction_log.txt (project root — not data/) this week. "
        "2-4 entries. Specific behaviors in specific contexts. "
        "Blank weeks score everything MEDIUM next cycle — the cross-reference gate loses precision.",
        styles["normal"]))

    story.append(Paragraph("Quarterly Eval", styles["h2"]))
    story.append(Paragraph(
        "If quarterly eval is due: feed change_log.json and friction_log.txt to Claude. "
        "Request eval report. Revert changes that did not hold. "
        "Flag source signal from reverted changes as low-quality in archive.",
        styles["normal"]))


def pdf_build_claire_a_section(story: list, decision_record: dict | None, styles: dict):
    """PDF Section 6 — CLAIRE-A Shadow Decisions."""
    story.append(Paragraph("CLAIRE-A Shadow Decisions", styles["h1"]))

    if decision_record is None:
        story.append(Paragraph(
            "No CLAIRE-A decision record found for this cycle. "
            "Run claire_a_assembler.py and claire_a_runner.py after synthesis.",
            styles["normal"]))
        story.append(pdf_rule(styles))
        return

    apply_count = decision_record.get("apply_count", 0)
    skip_count  = decision_record.get("skip_count", 0)
    defer_count = decision_record.get("defer_count", 0)
    total       = apply_count + skip_count + defer_count
    escalation  = decision_record.get("escalation_required", False)
    model       = decision_record.get("model", "unknown")

    story.append(Paragraph(
        f'<font color="#1E3A5F"><b>Engine evaluated {total} candidates ({model}): '
        f'{apply_count} apply &nbsp;|&nbsp; {skip_count} skip &nbsp;|&nbsp; '
        f'{defer_count} defer</b></font>',
        styles["normal"]))

    if escalation:
        reason = decision_record.get("escalation_reason", "")
        story.append(pdf_callout(f"ESCALATION REQUIRED: {reason}", styles))

    decisions = decision_record.get("decisions", [])
    if decisions:
        story.append(Paragraph("Candidate Decisions", styles["h2"]))

        headers = [
            Paragraph("Decision",   styles["table_header"]),
            Paragraph("Confidence", styles["table_header"]),
            Paragraph("Candidate",  styles["table_header"]),
            Paragraph("Risk Flags", styles["table_header"]),
        ]

        order = {"apply": 0, "defer": 1, "skip": 2}
        sorted_decisions = sorted(
            decisions, key=lambda d: order.get(d.get("decision", "skip"), 2))

        table_data = [headers]
        for row_idx, d in enumerate(sorted_decisions):
            decision = d.get("decision", "")
            conf     = d.get("confidence", 0.0)
            summary  = d.get("candidate_summary", "")[:120]
            flags    = ", ".join(d.get("risk_flags", [])) or "-"

            # Decision cell color
            if decision == "apply":
                dec_color = "#007000"
            elif decision == "skip":
                dec_color = "#990000"
            else:
                dec_color = "#B88600"

            dec_para  = Paragraph(
                f'<font color="{dec_color}"><b>{decision.upper()}</b></font>',
                styles["table_body"])
            conf_para = Paragraph(f"{conf:.0%}", styles["table_body"])
            sum_para  = Paragraph(summary, styles["table_body"])
            flag_para = Paragraph(flags, styles["table_body"])
            table_data.append([dec_para, conf_para, sum_para, flag_para])

        col_w = [1.0*inch, 0.8*inch, 3.5*inch, 1.2*inch]
        t = Table(table_data, colWidths=col_w)
        t.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0), PDF_NAVY),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1), [PDF_WHITE, PDF_LGRAY]),
            ("GRID",          (0, 0), (-1, -1), 0.5, PDF_BGRAY),
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING",   (0, 0), (-1, -1), 6),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
        ]))
        story.append(t)
        story.append(Spacer(1, 8))

    deferrals = [d for d in decisions if d.get("decision") == "defer"]
    if deferrals:
        story.append(Paragraph("Deferred — Conditions for Resolution", styles["h2"]))
        for d in deferrals:
            story.append(Paragraph(
                f'<b>{d.get("candidate_summary", "")[:100]}</b>',
                styles["body_bold"]))
            condition = d.get("defer_condition") or "No condition specified."
            story.append(pdf_callout(f"Resolve when: {condition}", styles))

    session_notes = decision_record.get("session_notes", "").strip()
    if session_notes:
        story.append(Paragraph("Engine Session Notes", styles["h2"]))
        story.append(Paragraph(session_notes, styles["normal"]))

    story.append(pdf_rule(styles))


def generate_techniques_pdf(date_str: str, technique_candidates: list) -> Path:
    """Build a lightweight Track C techniques PDF.

    Lighter than the main digest — no CLAIRE-A column, no triage stats.
    Sections: header, one entry per technique (name, description, source signal,
    confidence), footer with total count.

    Returns the output path.
    """
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = OUTPUT_DIR / f"claire_techniques_{ts}.pdf"

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=f"CLAIRE Technique Candidates {date_str}",
        author="CLAIRE",
        creator="CLAIRE",
    )

    styles = build_pdf_styles()
    story  = []

    # ── Header ────────────────────────────────────────────────────────────────
    story.append(Paragraph("CLAIRE — Technique Candidates", styles["title"]))
    story.append(Paragraph(f"Run date: {date_str}", styles["eyebrow"]))
    story.append(HRFlowable(width="100%", thickness=2, color=PDF_GOLD,
                            spaceAfter=12, spaceBefore=4))

    story.append(Paragraph(
        "Cross-platform workflow techniques surfaced this cycle. "
        "Test manually before considering any configuration change.",
        styles["normal"],
    ))
    story.append(Spacer(1, 8))

    # ── Technique entries ─────────────────────────────────────────────────────
    for i, technique in enumerate(technique_candidates, 1):
        name       = technique.get("technique_name", f"Technique {i}")
        description = technique.get("description", "")
        test_hint  = technique.get("test_suggestion", "")
        confidence = technique.get("confidence", "MEDIUM")
        sources    = technique.get("source_posts", [])
        source_str = ", ".join(sources[:3]) if sources else "—"

        conf_label = pdf_confidence_label(confidence)

        keep_items = []
        keep_items.append(Paragraph(
            f'<b>{i}. {name}</b>{conf_label}',
            styles["body_bold"],
        ))
        keep_items.append(Paragraph(description, styles["normal"]))
        keep_items.append(Paragraph(
            f'<b>Source signal:</b> {source_str}',
            styles["small"],
        ))
        if test_hint:
            keep_items.append(pdf_callout(f"Test: {test_hint}", styles))

        story.append(KeepTogether(keep_items))
        story.append(Spacer(1, 10))

    # ── Footer ────────────────────────────────────────────────────────────────
    story.append(HRFlowable(width="100%", thickness=1, color=PDF_BGRAY,
                            spaceAfter=6, spaceBefore=6))
    story.append(Paragraph(
        f"Total technique candidates this cycle: {len(technique_candidates)}",
        styles["small"],
    ))

    doc.build(story)
    log.info(f"Track C techniques PDF saved → {output_path.name}")
    return output_path


def generate_pdf(
    date_str: str,
    candidates: dict,
    triage_data: dict,
    claire_a_dr: dict | None,
) -> Path:
    """Build the six-section CLAIRE digest as a PDF and return the output path."""
    output_path = OUTPUT_DIR / f"claire_digest_{date_str}.pdf"

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=f"CLAIRE Digest {date_str}",
        author="CLAIRE",
        creator="CLAIRE",
        producer="",
        subject="",
    )

    styles = build_pdf_styles()
    story  = []

    # Title block
    story.append(Paragraph("Weekly Intelligence Digest", styles["eyebrow"]))
    story.append(Paragraph("CLAIRE", styles["title"]))
    story.append(Paragraph(
        f"Claude Configuration Signal Report — {date_str}", styles["subtitle"]))
    story.append(HRFlowable(width="100%", thickness=3, color=PDF_GOLD,
                            spaceAfter=16, spaceBefore=0))

    # Six sections
    pdf_build_signal_summary(story, triage_data, styles)
    pdf_build_track_a_section(story, candidates.get("a", {}), styles)
    pdf_build_track_b_section(story, candidates.get("b", {}), styles)
    pdf_build_track_c_section(story, candidates.get("c", {}), styles)
    pdf_build_eval_section(story, date_str, styles)
    pdf_build_claire_a_section(story, claire_a_dr, styles)

    doc.build(story)
    log.info(f"PDF digest saved → {output_path}")
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def load_latest_claire_a_decisions() -> dict | None:
    """Finds and loads the most recent CLAIRE-A decision record."""
    files = sorted(CLAIRE_A_DATA_DIR.glob("claire_a_decisions_*.json"), reverse=True)
    if not files:
        return None
    with open(files[0], encoding="utf-8") as f:
        wrapper = json.load(f)
    dr = wrapper.get("decision_record")
    if not dr:
        return None
    log.info(f"CLAIRE-A decisions loaded from {files[0].name}")
    return dr


def build_claire_a_section(doc: "Document", decision_record: dict | None):
    """Section 6 — CLAIRE-A Shadow Decisions."""
    doc.add_heading("CLAIRE-A Shadow Decisions", level=1)

    if decision_record is None:
        doc.add_paragraph(
            "No CLAIRE-A decision record found for this cycle. "
            "Run claire_a_assembler.py and claire_a_runner.py after synthesis."
        )
        add_rule(doc)
        return

    # Summary line
    apply_count  = decision_record.get("apply_count", 0)
    skip_count   = decision_record.get("skip_count", 0)
    defer_count  = decision_record.get("defer_count", 0)
    total        = apply_count + skip_count + defer_count
    escalation   = decision_record.get("escalation_required", False)
    model        = decision_record.get("model", "unknown")

    summary_para = doc.add_paragraph()
    summary_para.paragraph_format.space_after = Pt(8)
    run = summary_para.add_run(
        f"Engine evaluated {total} candidates ({model}): "
        f"{apply_count} apply  |  {skip_count} skip  |  {defer_count} defer"
    )
    run.font.name  = "Century Gothic"
    run.font.size  = Pt(11)
    run.font.color.rgb = NAVY

    if escalation:
        reason = decision_record.get("escalation_reason", "")
        add_callout(doc, f"ESCALATION REQUIRED: {reason}")

    # Decision table
    decisions = decision_record.get("decisions", [])
    if decisions:
        doc.add_heading("Candidate Decisions", level=2)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["Decision", "Confidence", "Candidate", "Risk Flags"]
        for i, h in enumerate(headers):
            header_cell(table.rows[0].cells[i], h)

        # Sort: apply first, then defer, then skip
        order = {"apply": 0, "defer": 1, "skip": 2}
        sorted_decisions = sorted(
            decisions, key=lambda d: order.get(d.get("decision", "skip"), 2)
        )

        for row_idx, d in enumerate(sorted_decisions):
            row      = table.add_row()
            decision = d.get("decision", "")
            conf     = d.get("confidence", 0.0)
            summary  = d.get("candidate_summary", "")[:120]
            flags    = ", ".join(d.get("risk_flags", [])) or "-"
            even     = row_idx % 2 == 0

            # Decision cell — colour-coded
            decision_cell = row.cells[0]
            body_cell(decision_cell, decision.upper(), even)
            if decision == "apply":
                decision_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x70, 0x00)
            elif decision == "skip":
                decision_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x99, 0x00, 0x00)
            elif decision == "defer":
                decision_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xB8, 0x86, 0x00)

            body_cell(row.cells[1], f"{conf:.0%}", even)
            body_cell(row.cells[2], summary, even)
            body_cell(row.cells[3], flags, even)

        doc.add_paragraph()  # spacing after table

    # Deferrals — expanded detail
    deferrals = [d for d in decisions if d.get("decision") == "defer"]
    if deferrals:
        doc.add_heading("Deferred — Conditions for Resolution", level=2)
        for d in deferrals:
            para = doc.add_paragraph()
            run  = para.add_run(d.get("candidate_summary", "")[:100])
            run.font.bold      = True
            run.font.name      = "Century Gothic"
            run.font.color.rgb = NAVY
            condition = d.get("defer_condition") or "No condition specified."
            add_callout(doc, f"Resolve when: {condition}")

    # Session notes — cross-batch observations
    session_notes = decision_record.get("session_notes", "").strip()
    if session_notes:
        doc.add_heading("Engine Session Notes", level=2)
        doc.add_paragraph(session_notes)

    add_rule(doc)


def main():
    parser = argparse.ArgumentParser(description="CLAIRE output layer")
    parser.add_argument("--date",   help="Override digest date (YYYY-MM-DD)")
    parser.add_argument(
        "--format",
        choices=["docx", "pdf"],
        default="docx",
        help="Output format (default: docx)",
    )
    args = parser.parse_args()

    date_str  = args.date or datetime.now(timezone.utc).strftime("%Y-%m-%d")
    run_start = datetime.now(timezone.utc)

    log.info(f"CLAIRE output started — date={date_str}  format={args.format}")

    # Load candidate files
    candidates = {}
    for track, path in CANDIDATE_PATHS.items():
        if path.exists():
            with open(path, encoding="utf-8") as f:
                data = json.load(f)
            candidates[track] = data.get("candidates", {})
            log.info(f"Track {track.upper()} candidates loaded from {path.name}")
        else:
            log.warning(f"Candidate file missing: {path.name} — section will show empty")
            candidates[track] = {}

    # Load triage metadata for signal summary
    triage_data = {}
    if TRIAGE_TAGGED_PATH.exists():
        with open(TRIAGE_TAGGED_PATH, encoding="utf-8") as f:
            triage_data = json.load(f)

    # CLAIRE-A shadow decisions
    claire_a_dr = load_latest_claire_a_decisions()

    if args.format == "pdf":
        # ── PDF path ─────────────────────────────────────────────────────────
        try:
            output_path = generate_pdf(date_str, candidates, triage_data, claire_a_dr)
        except Exception as e:
            log.error(f"PDF generation failed: {e}")
            sys.exit(1)

        log.info("═" * 60)
        log.info("CLAIRE output complete.")
        log.info(f"PDF digest: {output_path}")
        # Emit digest_date for GitHub Actions step output
        print(f"digest_date={date_str}", flush=True)

    else:
        # ── DOCX path (default) ───────────────────────────────────────────────
        doc = init_document()
        add_title_block(doc, date_str)

        build_signal_summary(doc, triage_data)
        build_track_a_section(doc, candidates.get("a", {}))
        build_track_b_section(doc, candidates.get("b", {}))
        build_track_c_section(doc, candidates.get("c", {}))
        build_eval_section(doc, date_str)
        build_claire_a_section(doc, claire_a_dr)

        filename = f"CLAIRE_Weekly_Digest_{date_str}.docx"
        output_path = OUTPUT_DIR / filename
        try:
            doc.save(str(output_path))
            log.info(f"Digest saved → {output_path}")
        except PermissionError:
            locked_path = OUTPUT_DIR / ("CLAIRE_Weekly_Digest_" + date_str + "_locked.docx")
            doc.save(str(locked_path))
            log.warning("Output file was locked - saved to fallback: " + locked_path.name)
        except Exception as e:
            log.error("Failed to save digest: " + str(e))
            sys.exit(1)

        skill_count = write_skill_drafts(candidates.get("b", {}))
        log.info("Skill drafts written: " + str(skill_count))

        log.info("=" * 60)
        log.info("CLAIRE output complete.")
        log.info("Digest: " + str(output_path))
        log.info("Skill drafts: " + str(skill_count) + " files in skill_drafts/")
        log.info("Review the digest. Apply candidates manually with hypotheses.")

    # ── Build 8: Track C separate techniques PDF ──────────────────────────────
    # Check synthesis_queue_track_c.json for Track C activity this cycle.
    # PDF content sourced from candidates_track_c.json (synthesized output).
    track_c_queue_path = DATA_DIR / "synthesis_queue_track_c.json"
    track_c_has_posts  = False
    if track_c_queue_path.exists():
        try:
            with open(track_c_queue_path, encoding="utf-8") as f:
                tc_queue = json.load(f)
            track_c_has_posts = len(tc_queue.get("posts", [])) > 0
        except (json.JSONDecodeError, OSError) as e:
            log.warning(f"Could not read synthesis_queue_track_c.json: {e}")

    technique_candidates = candidates.get("c", {}).get("technique_candidates", [])

    if track_c_has_posts and technique_candidates:
        try:
            techniques_path = generate_techniques_pdf(date_str, technique_candidates)
            log.info(f"Track C techniques PDF: {techniques_path.name}")
        except Exception as e:
            log.error(f"Track C techniques PDF generation failed: {e} — continuing")
    else:
        log.info("Track C: no candidates this cycle — techniques PDF skipped")
    # ─────────────────────────────────────────────────────────────────────────


if __name__ == "__main__":
    main()
